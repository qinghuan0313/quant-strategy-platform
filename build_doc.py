# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

for s in doc.sections:
    s.top_margin = Cm(2.54)
    s.bottom_margin = Cm(2.54)
    s.left_margin = Cm(3.17)
    s.right_margin = Cm(3.17)

sty = doc.styles["Normal"]
sty.font.name = "Times New Roman"
sty.font.size = Pt(12)
sty.paragraph_format.line_spacing = 1.5
sty.paragraph_format.space_after = Pt(0)
sty.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

for lv in range(1, 4):
    hs = doc.styles["Heading %d" % lv]
    hs.font.name = "Times New Roman"
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.font.bold = True
    hs.paragraph_format.line_spacing = 1.5
    hs.paragraph_format.space_before = Pt(12)
    hs.paragraph_format.space_after = Pt(6)
    if lv == 1:
        hs.font.size = Pt(16)
    elif lv == 2:
        hs.font.size = Pt(14)
    else:
        hs.font.size = Pt(12)


def para(text, bold=False, fs=12, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    r.font.size = Pt(fs)
    r.bold = bold


def code(text):
    for line in text.strip().split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.2
        r = p.add_run(line)
        r.font.name = "Consolas"
        r.font.size = Pt(9)


def cell_font(cell, text, fn_cn, fs, bold=False, align="left"):
    cell.text = ""
    p = cell.paragraphs[0]
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.element.rPr.rFonts.set(qn('w:eastAsia'), fn_cn)
    r.font.size = Pt(fs)
    r.bold = bold


def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    for i, h in enumerate(headers):
        cell_font(t.rows[0].cells[i], h, "黑体", 9, True, "center")
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            cell_font(t.rows[ri + 1].cells[ci], str(v), "宋体", 9)
    doc.add_paragraph()


# =============================
#   FIX: strip theme fonts from styles, use explicit font names
# =============================
def fix_styles(doc):
    """Remove theme font references from styles so explicit fonts take effect."""
    from lxml import etree
    nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    for part in doc.part.related_parts.values():
        if hasattr(part, 'element') and part.element.tag.endswith('}styles'):
            root = part.element
            # Find all rFonts elements
            for rFonts in root.iter('{%s}rFonts' % nsmap['w']):
                # Remove theme font attributes that override explicit fonts
                for attr in ['{%s}eastAsiaTheme' % nsmap['w'],
                             '{%s}asciiTheme' % nsmap['w'],
                             '{%s}hAnsiTheme' % nsmap['w'],
                             '{%s}cstheme' % nsmap['w']]:
                    if attr in rFonts.attrib:
                        del rFonts.attrib[attr]


# ============ BUILD ============
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(72)
p.paragraph_format.space_after = Pt(6)
r = p.add_run("量化学习 · 阶段一复习文档")
r.font.name = "Times New Roman"
r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
r.font.size = Pt(22)
r.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Python x 量化数据基础")
r.font.name = "Times New Roman"
r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
r.font.size = Pt(15)

para("学习时间：2026年6月11日 — 6月18日", indent=False)
para("数据素材：沪深300指数（akshare, sh000300）", indent=False)
para("学习原则：每行代码先解释再运行；编程为量化服务；用到什么学什么；不背语法，理解逻辑", indent=False)
doc.add_page_break()

# 1.1
doc.add_heading("1.1 变量和数据类型", 1)
add_table(
    ["概念", "说明", "量化场景"],
    [
        ["变量", "存数据的盒子，= 是把右边装进左边", "price_today = 100"],
        ["数字(int/float)", "整数和小数", "price = 3500.5"],
        ["字符串(str)", "文字，用引号包裹", '"CSI 300"'],
        ["列表(list)", "方括号装一组数据", "[100, 102, 99, 105, 98]"],
        ["列表取值", "prices[0]取第1个, Python从0数起", "prices[0] = 第1天价格"],
    ],
)
doc.add_heading("代码", 2)
code("""price_today = 100           # 把100存进变量
price_yesterday = 98        # 把98存进变量
return_rate = (price_today - price_yesterday) / price_yesterday
print(return_rate)          # 0.0204 -> 涨了2.04%""")

# 1.2
doc.add_heading("1.2 for 循环", 1)
add_table(
    ["概念", "说明", "例子"],
    [
        ["for i in range(n)", "用位置编号循环, i = 0,1,2...n-1", "range(4) -> 0,1,2,3"],
        ["for item in list", "直接用值循环, 每轮拿到一个元素", "for price in prices:"],
        ["len(list)", "返回列表元素个数", "len([100,102,99]) = 3"],
    ],
)
doc.add_heading("代码：用 for 循环算所有收益率", 2)
code("""prices = [100, 102, 99, 105, 98]
returns = []

for i in range(len(prices) - 1):
    ret = (prices[i+1] - prices[i]) / prices[i]
    returns.append(ret)

print(returns)
# [0.02, -0.0294, 0.0606, -0.0667]""")
doc.add_heading("循环执行过程（4趟）", 2)
add_table(
    ["第几趟", "i", "计算表达式", "结果"],
    [
        ["1", "0", "(102-100)/100", "+2.00%"],
        ["2", "1", "(99-102)/102", "-2.94%"],
        ["3", "2", "(105-99)/99", "+6.06%"],
        ["4", "3", "(98-105)/105", "-6.67%"],
    ],
)

# 1.3
doc.add_heading("1.3 if-elif-else 判断", 1)
add_table(
    ["关键字", "作用", "什么时候执行"],
    [
        ["if 条件", "第一个判断", "条件为真时"],
        ["elif 条件", "第二个判断 (else+if)", "前面都不满足且本条件为真时"],
        ["else", "兜底", "前面所有条件都不满足时"],
    ],
)
doc.add_heading("代码：判断涨跌平", 2)
code("""prices = [100, 102, 99, 105, 98, 103, 101]

for i in range(len(prices) - 1):
    today = prices[i + 1]
    yesterday = prices[i]

    if today > yesterday:
        print(f"第{i+2}天：涨了")
    elif today < yesterday:
        print(f"第{i+2}天：跌了")
    else:
        print(f"第{i+2}天：平了")""")
para("关键注意：", bold=True)
para("(1) print 必须在 for 循环内（缩进4个空格），放在循环外只会打印最后一次结果。")
para('(2) f"..." 是格式化字符串，花括号 {变量} 会被替换成变量的值。')
para("(3) {ret:.2%} 可自动把小数转成百分比并保留两位小数。")

# 1.4
doc.add_heading("1.4 函数", 1)
add_table(
    ["关键字", "作用", "例子"],
    [
        ["def", "定义（发明）一个新工具", "def add(a, b):"],
        ["参数", "工具接收的输入值", "(a, b) 就是参数"],
        ["return", "把计算结果交出去给外面用", "return a + b"],
        ["调用", "使用函数，传入具体数值", "result = add(3, 5)"],
    ],
)
doc.add_heading("print 与 return 的区别", 2)
add_table(
    ["", "print", "return"],
    [
        ["作用", "显示在屏幕上给人看", "把值交出去给别的代码用"],
        ["值能否复用", "否，显示完就没了", "是，可以存进变量继续用"],
        ["用在何处", "想看结果的时候", "函数内部；函数外用 print 打印返回值"],
    ],
)
doc.add_heading("代码：收益率计算函数", 2)
code("""def calc_return(today_price, yesterday_price):
    ret = (today_price - yesterday_price) / yesterday_price
    return ret

result = calc_return(105, 100)
print(result)   # 0.05""")
doc.add_heading("代码：市场判断函数（复习题第2题）", 2)
code("""def judge_market(r):
    if r > 0.01:
        return "大涨"
    elif r > 0:
        return "小涨"
    elif r > -0.01:
        return "小跌"
    else:
        return "大跌"

print(judge_market(0.02))     # 大涨
print(judge_market(-0.005))   # 小跌
print(judge_market(0.003))    # 小涨""")
para("注意事项：", bold=True)
para("(1) 1% 要写成 0.01，Python 不支持百分号写法。")
para("(2) 函数内部用 return 不用 print；在外面用 print 打印函数的返回值。")
para("(3) return 不是函数，习惯上不加括号，直接 return 值。")

# 1.5
doc.add_heading("1.5 pandas 读数据", 1)
add_table(
    ["概念", "说明", "代码"],
    [
        ["导入 pandas", "导入并起短名 pd", "import pandas as pd"],
        ["pd.DataFrame()", "把字典转成表格", "df = pd.DataFrame(data)"],
        ['df["列名"]', "取某一列", 'df["close"]'],
        [".head(n)", "看前 n 行，默认5", "df.head(10)"],
        [".tail(n)", "看后 n 行", "df.tail()"],
        [".describe()", "快速统计", "df.describe()"],
        [".columns", "查看有哪些列", "df.columns"],
        ["len(df)", "总行数", "len(df)"],
        ["type(x)", "查看变量是什么类型", "type(df)"],
    ],
)
doc.add_heading("代码：手写 DataFrame", 2)
code("""import pandas as pd

data = {
    "日期": ["2020-01-02", "2020-01-03", "2020-01-06"],
    "收盘价": [100, 102, 99]
}
df = pd.DataFrame(data)
print(df)""")
doc.add_heading("代码：拉真实沪深300数据", 2)
code("""import akshare as ak

data = ak.stock_zh_index_daily(symbol="sh000300")
print(type(data))       # <class 'pandas.DataFrame'>
print(data.columns)     # date, open, high, low, close, volume
print(len(data))        # ~5930 个交易日
print(data.head())      # 看前5行""")

# 1.6
doc.add_heading("1.6 pandas 做计算", 1)
add_table(
    ["方法", "作用", "注意"],
    [
        [".pct_change()", "算(今天-昨天)/昨天", "第一行是 NaN（没有前一天）"],
        [".mean()", "求平均值", "自动忽略 NaN"],
        [".sum()", "求和", "True=1, False=0 -> 自动计数"],
        ["(条件).sum()", "数满足条件的有多少行", '(data["close"] > 5000).sum()'],
    ],
)
doc.add_heading("代码", 2)
code("""ret = data["close"].pct_change()
print(ret.head(10))       # 第一行 NaN

print(ret.mean())          # ~0.00034 (日赚0.034%, 年化~9%)
print((data["close"] > 5000).sum())   # 250天""")
para("易错提醒：pct_change 是函数，必须加括号 pct_change()；.head()、.describe()、.mean()、.sum() 同理。")

# ============ REVIEW ============
doc.add_page_break()
doc.add_heading("阶段一复习题（三题全部通过）", 1)

doc.add_heading("第1题：变量 + 列表 + for 循环", 2)
para("题目：给定沪深300连续5天收盘价，用 for 循环算出每天比第一天涨了多少，格式化输出百分比。")
code("""close = [3500, 3520, 3480, 3550, 3530]

for i in range(len(close)):
    ret = (close[i] - close[0]) / close[0]
    print(f"第{i+1}天: {ret:.2%}")
# 第1天: 0.00% / 第2天: 0.57% / 第3天: -0.57%
# 第4天: 1.43% / 第5天: 0.86%""")
para("易错点：print 要缩进在 for 循环里面；用 len(close) 而非 len(close)-1。")

doc.add_heading("第2题：if 判断 + 函数 + return", 2)
para("题目：写 judge_market 函数，根据日收益率返回大涨/小涨/小跌/大跌。")
code("""def judge_market(r):
    if r > 0.01:
        return "大涨"
    elif r > 0:
        return "小涨"
    elif r > -0.01:
        return "小跌"
    else:
        return "大跌"

print(judge_market(0.02))     # 大涨
print(judge_market(-0.005))   # 小跌
print(judge_market(0.003))    # 小涨""")
para("易错点：1% 要写成 0.01 不能写 1%；函数内用 return 不用 print。")

doc.add_heading("第3题：pandas 综合应用", 2)
para("题目：拉沪深300数据，.describe() 看统计，算日均收益率，统计收盘价>5000的天数。")
code("""import akshare as ak

data = ak.stock_zh_index_daily(symbol="sh000300")
print(data.describe())

ret = data["close"].pct_change()
print("日均收益率:", ret.mean())                         # ~0.00034
print("收盘价>5000的天数:", (data["close"] > 5000).sum())  # 250""")
para("易错点：pct_change 后面必须加 ()；条件表达式要用括号括起来再 .sum()。")

# ============ CHEAT SHEET ============
doc.add_page_break()
doc.add_heading("附录：Python 符号速查表", 1)
add_table(
    ["符号", "名称", "用途", "例子"],
    [
        ["()", "圆括号", "调用函数，做动作", "print(), len(), .mean()"],
        ["[]", "方括号", "取列表元素 / 创建列表", 'prices[0], df["close"], [1,2,3]'],
        ['"" / \'\'', "引号", "包裹字符串文字", '"沪深300"'],
        ["{}", "花括号", "创建字典；f-string 嵌入变量", '{"key": val}, f"第{i}天"'],
        [":", "冒号", "代码块开头标记", "def, if, elif, else, for 后面加"],
        ["#", "井号", "注释，Python 不执行", "# 这是一行注释"],
        ["=", "单等号", "赋值：把右边装进左边", "price = 100"],
        ["==", "双等号", "判断是否相等", "if a == b:"],
        [".", "点号", "的：访问内部功能", 'df["close"].mean()'],
        ["    ", "缩进(4空格)", "表示代码属于哪个代码块", "for/if/def 内代码要缩进"],
    ],
)

# ============ FIX AND SAVE ============
fix_styles(doc)

save_path = r"C:\Users\szh\Desktop\quant_learn\阶段一复习.docx"
doc.save(save_path)
print("Saved: " + save_path)
