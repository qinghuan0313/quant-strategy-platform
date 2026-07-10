# -*- coding: utf-8 -*-
"""
Generate Phase 1 review docx - university thesis formatting
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# Normal style
style = doc.styles["Normal"]
style.font.name = "SimSun"
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

# Heading styles
for lvl in range(1, 4):
    hs = doc.styles["Heading %d" % lvl]
    hs.font.name = "SimHei"
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.font.bold = True
    hs.paragraph_format.line_spacing = 1.5
    hs.paragraph_format.space_before = Pt(12)
    hs.paragraph_format.space_after = Pt(6)
    if lvl == 1:
        hs.font.size = Pt(16)
    elif lvl == 2:
        hs.font.size = Pt(14)
    else:
        hs.font.size = Pt(12)


def add_para(text, bold=False, font_size=12, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(text)
    run.font.name = "SimSun"
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.size = Pt(font_size)
    run.bold = bold
    return p


def add_code_block(code_text):
    for line in code_text.strip().split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.2
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def set_cell_font(cell, text, font_name, font_size, bold=False, align="left"):
    cell.text = ""
    p = cell.paragraphs[0]
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = font_name
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.bold = bold


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
        '<w:tblPr %s></w:tblPr>' % nsdecls("w"))
    borders = parse_xml(
        '<w:tblBorders %s>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>' % nsdecls("w")
    )
    tblPr.append(borders)

    for i, h in enumerate(headers):
        set_cell_font(table.rows[0].cells[i], h, "SimHei", 9, bold=True, align="center")

    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            set_cell_font(table.rows[r + 1].cells[c], str(val), "SimSun", 9)

    doc.add_paragraph()
    return table


# ==================== COVER ====================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(72)
p.paragraph_format.space_after = Pt(6)
run = p.add_run("量化学习 · 阶段一复习文档")
run.font.name = "SimHei"
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
run.font.size = Pt(22)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(24)
run = p.add_run("Python × 量化数据基础")
run.font.name = "SimHei"
run.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
run.font.size = Pt(16)

add_para("Study period: June 11-18, 2026", indent=False)
add_para("Data: CSI 300 Index via akshare (sh000300)", indent=False)
add_para("Principle: explain every line before running; programming serves quant; learn what you use; understand logic, not syntax", indent=False)

doc.add_page_break()

# ==================== 1.1 ====================
doc.add_heading("1.1 Variables and Data Types", level=1)
doc.add_heading("Key Concepts", level=2)

add_table(
    ["Concept", "Explanation", "Quant Example"],
    [
        ["variable", "a box that stores data; = means assign right to left", "price_today = 100"],
        ["number (int/float)", "integer and decimal", "price = 3500.5"],
        ["string (str)", "text wrapped in quotes", '"CSI 300"'],
        ["list", "square brackets hold a group of data", "[100, 102, 99, 105, 98]"],
        ["list indexing", "prices[0] gets 1st element; Python counts from 0", "prices[0] = day 1 price"],
    ],
)

doc.add_heading("Code", level=2)
add_code_block("""price_today = 100
price_yesterday = 98
return_rate = (price_today - price_yesterday) / price_yesterday
print(return_rate)   # 0.0204 -> up 2.04%""")

# ==================== 1.2 ====================
doc.add_heading("1.2 for Loop", level=1)
doc.add_heading("Key Concepts", level=2)

add_table(
    ["Concept", "Explanation", "Example"],
    [
        ["for i in range(n)", "loop with index i = 0,1,2...n-1", "range(4) -> 0,1,2,3"],
        ["for item in list", "loop directly over values", "for price in prices:"],
        ["len(list)", "return number of elements", "len([100,102,99]) = 3"],
    ],
)

doc.add_heading("Code: Calculate All Returns with for Loop", level=2)
add_code_block("""prices = [100, 102, 99, 105, 98]
returns = []                        # empty list for results

for i in range(len(prices) - 1):    # 5 prices -> 4 return segments
    ret = (prices[i+1] - prices[i]) / prices[i]
    returns.append(ret)             # append to list

print(returns)
# [0.02, -0.0294, 0.0606, -0.0667]""")

doc.add_heading("Loop Execution", level=2)
add_table(
    ["Iteration", "i", "Expression", "Result"],
    [
        ["1", "0", "(102-100)/100", "+2.00%"],
        ["2", "1", "(99-102)/102", "-2.94%"],
        ["3", "2", "(105-99)/99", "+6.06%"],
        ["4", "3", "(98-105)/105", "-6.67%"],
    ],
)

# ==================== 1.3 ====================
doc.add_heading("1.3 if-elif-else Conditionals", level=1)
doc.add_heading("Key Concepts", level=2)

add_table(
    ["Keyword", "Purpose", "When It Executes"],
    [
        ["if cond", "first check", "when condition is True"],
        ["elif cond", "second check (else + if)", "when all above are False and this is True"],
        ["else", "catch-all", "when ALL above conditions are False"],
    ],
)

doc.add_heading("Code: Detect Up/Down/Flat Days", level=2)
add_code_block("""prices = [100, 102, 99, 105, 98, 103, 101]

for i in range(len(prices) - 1):
    today = prices[i + 1]
    yesterday = prices[i]

    if today > yesterday:
        print(f"Day {i+2}: UP")
    elif today < yesterday:
        print(f"Day {i+2}: DOWN")
    else:
        print(f"Day {i+2}: FLAT")""")

add_para("Key notes:", bold=True)
add_para("(1) print must be indented inside the for loop (4 spaces); outside = only last result printed.")
add_para('(2) f"..." is an f-string: {var} gets replaced by the variable value.')
add_para("(3) {ret:.2%} auto-formats decimal as percentage with 2 decimal places.")

# ==================== 1.4 ====================
doc.add_heading("1.4 Functions", level=1)
doc.add_heading("Key Concepts", level=2)

add_table(
    ["Keyword", "Purpose", "Example"],
    [
        ["def", "define (create) a new tool", "def add(a, b):"],
        ["parameters", "input values the tool accepts", "(a, b) are parameters"],
        ["return", "hand the result back to the caller", "return a + b"],
        ["call", "use the function with concrete values", "result = add(3, 5)"],
    ],
)

doc.add_heading("print vs return", level=2)
add_table(
    ["", "print", "return"],
    [
        ["Purpose", "show on screen for human eyes", "hand value back to other code"],
        ["Reusable?", "no, gone after display", "yes, can be stored in variable"],
        ["Where", "when you want to see output", "inside function; outside use print()"],
    ],
)

doc.add_heading("Code: Return Calculator Function", level=2)
add_code_block("""def calc_return(today_price, yesterday_price):
    ret = (today_price - yesterday_price) / yesterday_price
    return ret

result = calc_return(105, 100)
print(result)   # 0.05""")

doc.add_heading("Code: Market Judge Function (Review Q2)", level=2)
add_code_block("""def judge_market(r):
    if r > 0.01:       # > 1% -> big up
        return "big up"
    elif r > 0:        # 0 ~ 1% -> small up
        return "small up"
    elif r > -0.01:    # -1% ~ 0 -> small down
        return "small down"
    else:              # < -1% -> big down
        return "big down"

print(judge_market(0.02))     # big up
print(judge_market(-0.005))   # small down
print(judge_market(0.003))    # small up""")

add_para("Notes:", bold=True)
add_para("(1) 1% must be written as 0.01; Python does not support percent notation.")
add_para("(2) Use return inside function, print outside to display the return value.")
add_para("(3) return is not a function; idiom is return 'value' without parentheses.")

# ==================== 1.5 ====================
doc.add_heading("1.5 pandas: Reading Data", level=1)
doc.add_heading("Key Concepts", level=2)

add_table(
    ["Concept", "Explanation", "Code"],
    [
        ["import pandas", "import and alias as pd", "import pandas as pd"],
        ["pd.DataFrame()", "create table from dict", "df = pd.DataFrame(data)"],
        ["df[col]", "select one column", 'df["close"]'],
        [".head(n)", "view first n rows (default 5)", "df.head(10)"],
        [".tail(n)", "view last n rows", "df.tail()"],
        [".describe()", "quick stats: mean/std/min/max...", "df.describe()"],
        [".columns", "list column names", "df.columns"],
        ["len(df)", "total row count", "len(df)"],
        ["type(x)", "check variable type", "type(df)"],
    ],
)

doc.add_heading("Code: Manual DataFrame", level=2)
add_code_block("""import pandas as pd

data = {
    "date": ["2020-01-02", "2020-01-03", "2020-01-06"],
    "close": [100, 102, 99]
}

df = pd.DataFrame(data)
print(df)""")

doc.add_heading("Code: Pull Real CSI 300 Data", level=2)
add_code_block("""import akshare as ak

data = ak.stock_zh_index_daily(symbol="sh000300")
print(type(data))       # <class 'pandas.DataFrame'>
print(data.columns)     # ['date','open','high','low','close','volume']
print(len(data))        # ~5930 rows
print(data.head())      # first 5 rows""")

# ==================== 1.6 ====================
doc.add_heading("1.6 pandas: Computation", level=1)
doc.add_heading("Key Concepts", level=2)

add_table(
    ["Method", "Purpose", "Note"],
    [
        [".pct_change()", "compute (today-yesterday)/yesterday for entire column", "first row is NaN (no prior day)"],
        [".mean()", "compute average", "automatically ignores NaN"],
        [".sum()", "sum values; with condition counts rows", "True=1, False=0"],
        ["(cond).sum()", "count rows matching condition", '(df["close"] > 5000).sum()'],
    ],
)

doc.add_heading("Code", level=2)
add_code_block("""ret = data["close"].pct_change()
print(ret.head(10))       # first is NaN

# average daily return
print(ret.mean())          # ~0.00034 (0.034% daily, ~9% annualized)

# days where close > 5000 (2007 and 2015 bull markets)
print((data["close"] > 5000).sum())   # 250 days""")

add_para("Common mistake: pct_change is a FUNCTION; must write pct_change() with parentheses. Same for .head(), .describe(), .mean(), .sum(). Skipping () silently does nothing.")

# ==================== REVIEW ====================
doc.add_page_break()
doc.add_heading("Phase 1 Review Exercises (All Passed)", level=1)

doc.add_heading("Q1: Variables + List + for Loop", level=2)
add_para("Task: given 5 days of CSI 300 closes, compute each day's return vs day 1, format as percentage.")
add_code_block("""close = [3500, 3520, 3480, 3550, 3530]

for i in range(len(close)):
    ret = (close[i] - close[0]) / close[0]
    print(f"Day {i+1}: {ret:.2%}")

# Day 1: 0.00%
# Day 2: 0.57%
# Day 3: -0.57%
# Day 4: 1.43%
# Day 5: 0.86%""")

doc.add_heading("Q2: if + Function + return", level=2)
add_para("Task: write judge_market(r) that returns big up / small up / small down / big down.")
add_code_block("""def judge_market(r):
    if r > 0.01:
        return "big up"
    elif r > 0:
        return "small up"
    elif r > -0.01:
        return "small down"
    else:
        return "big down"

print(judge_market(0.02))     # big up
print(judge_market(-0.005))   # small down
print(judge_market(0.003))    # small up""")

doc.add_heading("Q3: pandas Comprehensive", level=2)
add_para("Task: pull CSI 300, describe close, compute avg daily return, count days where close > 5000.")
add_code_block("""import akshare as ak

data = ak.stock_zh_index_daily(symbol="sh000300")
print(data.describe())

ret = data["close"].pct_change()
print("Avg daily return:", ret.mean())                    # ~0.00034
print("Days close>5000:", (data["close"] > 5000).sum())   # 250""")

# ==================== CHEAT SHEET ====================
doc.add_page_break()
doc.add_heading("Appendix: Python Symbol Reference", level=1)

add_table(
    ["Symbol", "Name", "Usage", "Example"],
    [
        ["()", "parentheses", "call function / do action", "print(), len(), .mean(), .pct_change()"],
        ["[]", "brackets", "access list element / create list", 'prices[0], df["close"], [1,2,3]'],
        ['"" / \'\'', "quotes", "wrap string text", '"CSI 300"'],
        ["{}", "braces", "create dict / f-string interpolation", '{"k": v}, f"Day {i}"'],
        [":", "colon", "start code block", "after def, if, elif, else, for"],
        ["#", "hash", "comment, ignored by Python", "# this is a comment"],
        ["=", "equals", "assignment: put right into left", "price = 100"],
        ["==", "double equals", "equality check", "if a == b:"],
        [".", "dot", "'s — access internal method/attr", 'df["close"].mean() = mean of close column'],
    ],
)

save_path = r"C:\Users\szh\Desktop\quant_learn\Phase1_Review.docx"
doc.save(save_path)
print("Saved to: " + save_path)
