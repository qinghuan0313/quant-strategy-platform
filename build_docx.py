# -*- coding: utf-8 -*-
"""Build the final DOCX from content files and format specifications."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re, os, zipfile, shutil

# ── helpers ──────────────────────────────────────────────

def set_run_font(run, font_cn, font_en, size_pt, bold=False):
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.name = font_en
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml('<w:rFonts ' + nsdecls("w") + ' />')
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_cn)
    run.font.color.rgb = RGBColor(0, 0, 0)

def set_spacing(para, line_sp=1.5, before=0, after=0, indent=None):
    pf = para.paragraph_format
    pf.line_spacing = line_sp
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if indent is not None:
        pf.first_line_indent = Pt(indent)

def add_h1(doc, text):
    p = doc.add_paragraph()
    set_spacing(p, 1.5, 12, 6)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    set_run_font(r, "黑体", "Times New Roman", 16, True)
    return p

def add_h2(doc, text):
    p = doc.add_paragraph()
    set_spacing(p, 1.5, 12, 6)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    set_run_font(r, "黑体", "Times New Roman", 14, True)
    return p

def add_body(doc, text):
    p = doc.add_paragraph()
    set_spacing(p, 1.5, indent=24)
    r = p.add_run(text)
    set_run_font(r, "宋体", "Times New Roman", 12, False)
    return p

def add_cover_title(doc, text, size=22):
    p = doc.add_paragraph()
    set_spacing(p, 1.5, indent=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, "黑体", "Times New Roman", size, True)
    return p

def add_cover_info(doc, text):
    p = doc.add_paragraph()
    set_spacing(p, 1.5, indent=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, "宋体", "Times New Roman", 12, False)
    return p

def fix_themes(docx_path):
    tmp = os.path.join(os.path.dirname(docx_path), "_tmp_fix")
    if os.path.exists(tmp):
        shutil.rmtree(tmp)
    with zipfile.ZipFile(docx_path, "r") as z:
        z.extractall(tmp)
    sp = os.path.join(tmp, "word", "styles.xml")
    with open(sp, "r", encoding="utf-8") as f:
        c = f.read()
    for a in ["w:eastAsiaTheme", "w:asciiTheme", "w:hAnsiTheme", "w:cstheme"]:
        c = re.sub(r'\s*' + a + '="[^"]*"', "", c)
    with open(sp, "w", encoding="utf-8") as f:
        f.write(c)
    os.remove(docx_path)
    with zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as z:
        for root, _, files in os.walk(tmp):
            for fn in files:
                fp = os.path.join(root, fn)
                z.write(fp, os.path.relpath(fp, tmp))
    shutil.rmtree(tmp)

# ── read content files ───────────────────────────────────

base = os.path.dirname(os.path.abspath(__file__))

def read_q(filename):
    with open(os.path.join(base, filename), "r", encoding="utf-8") as f:
        return f.read()

q1 = read_q("q1.txt")
q2 = read_q("q2.txt")
q3 = read_q("q3.txt")
q4 = read_q("q4.txt")
q5 = read_q("q5.txt")

# ── build document ───────────────────────────────────────

doc = Document()
sec = doc.sections[0]
sec.top_margin = Cm(2.54)
sec.bottom_margin = Cm(2.54)
sec.left_margin = Cm(3.17)
sec.right_margin = Cm(3.17)

# cover
add_cover_title(doc, "大学语文期末考试复习资料")
doc.add_paragraph()
add_cover_title(doc, "城市文化·科幻文学·主体性重构·爱情叙事·脱口秀", 14)
doc.add_paragraph()
add_cover_info(doc, "厦门大学经济学院")
add_cover_info(doc, "2026年6月")
doc.add_page_break()

# process each question file
import json

for qfile in ["q1.txt", "q2.txt", "q3.txt", "q4.txt", "q5.txt"]:
    text = read_q(qfile)
    blocks = text.strip().split("\n===\n")
    for block in blocks:
        block = block.strip()
        if not block:
            continue
        lines = block.split("\n", 1)
        tag = lines[0].strip()
        content = lines[1].strip() if len(lines) > 1 else ""
        if tag == "H1":
            add_h1(doc, content)
        elif tag == "H2":
            add_h2(doc, content)
        elif tag == "BODY":
            add_body(doc, content)

# save
out = os.path.join(os.path.expanduser("~"), "Desktop", "大学语文期末考试复习资料.docx")
doc.save(out)
print("Saved:", out)
fix_themes(out)
print("Theme fix done. All complete!")
